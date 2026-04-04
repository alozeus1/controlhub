"""Governed server-side tool functions for Agent Service."""

import csv
import hashlib
import io
import json
import os
from datetime import datetime

from sqlalchemy import or_

from app.models import Asset, Deployment, Person
from app.services.google_workspace import publish_to_drive as google_publish_to_drive
from app.services.google_workspace import publish_to_sheet as google_publish_to_sheet
from app.storage.s3_artifacts import (
    generate_artifact_presigned_url,
    put_artifact_object,
    read_artifact_object,
)
from app.utils.people_rbac import get_person_for_user


SUPPORTED_OUTPUT_TYPES = {"csv", "xlsx", "docx", "md"}
SUPPORTED_MODULE_SCOPES = {"people", "assets", "deployments"}
SUPPORTED_DESTINATION_TYPES = {"download", "google_drive_folder", "google_sheet_range"}

LOCAL_ARTIFACT_BUCKET = "local-artifacts"
SENSITIVE_FIELD_TOKENS = {"secret", "password", "token", "key", "credential", "api_key"}


def _mask_sensitive_value(value):
    if value in (None, ""):
        return value
    return "***masked***"


def _sanitize_row(row):
    safe = {}
    for key, value in row.items():
        key_lower = str(key).lower()
        if any(token in key_lower for token in SENSITIVE_FIELD_TOKENS):
            safe[key] = _mask_sensitive_value(value)
        else:
            safe[key] = value
    return safe


def _people_base_query(filters):
    query = Person.query

    search = (filters or {}).get("search")
    if search:
        pattern = f"%{search}%"
        query = query.filter(
            or_(
                Person.first_name.ilike(pattern),
                Person.last_name.ilike(pattern),
                Person.email.ilike(pattern),
            )
        )

    for field in ("team", "department", "cohort"):
        value = (filters or {}).get(field)
        if value:
            query = query.filter(getattr(Person, field) == value)

    is_active = (filters or {}).get("is_active")
    if is_active is not None:
        query = query.filter(Person.is_active == bool(is_active))

    return query.order_by(Person.created_at.desc())


def query_people(actor, filters=None):
    """Query people rows with role-aware row-level access."""
    filters = filters or {}
    query = _people_base_query(filters)

    people = query.all()

    actor_person = get_person_for_user(actor.id) if actor else None
    rows = []

    for person in people:
        employment = person.active_employment

        # Row-level access rules
        if actor.role == "people_manager":
            if not actor_person or not employment or employment.manager_person_id != actor_person.id:
                continue
        elif actor.role == "mentor":
            if (
                not actor_person
                or not employment
                or employment.employment_type != "intern"
                or employment.mentor_person_id != actor_person.id
            ):
                continue

        # Employment-scoped filters
        if filters.get("employment_type") and (
            not employment or employment.employment_type != filters.get("employment_type")
        ):
            continue
        if filters.get("intern_track") and (
            not employment or employment.intern_track != filters.get("intern_track")
        ):
            continue
        if filters.get("status") and (
            not employment or employment.status != filters.get("status")
        ):
            continue
        if filters.get("manager_person_id") and (
            not employment or employment.manager_person_id != int(filters.get("manager_person_id"))
        ):
            continue

        row = {
            "person_id": person.id,
            "full_name": person.full_name,
            "email": person.email,
            "phone": person.phone,
            "team": person.team,
            "department": person.department,
            "cohort": person.cohort,
            "employment_type": employment.employment_type if employment else None,
            "intern_track": employment.intern_track if employment else None,
            "employment_status": employment.status if employment else None,
            "title": employment.title if employment else None,
            "manager_name": employment.manager.full_name if (employment and employment.manager) else None,
            "mentor_name": employment.mentor.full_name if (employment and employment.mentor) else None,
            "start_date": employment.start_date.isoformat() if (employment and employment.start_date) else None,
            "end_date": employment.end_date.isoformat() if (employment and employment.end_date) else None,
            "is_active": person.is_active,
        }
        rows.append(_sanitize_row(row))

    limit = int(filters.get("limit", 1000))
    return rows[: max(1, min(limit, 10000))]


def query_assets(actor, filters=None):
    """Query assets rows (basic implementation)."""
    filters = filters or {}
    query = Asset.query.order_by(Asset.created_at.desc())

    for field in ("asset_type", "status", "team", "department"):
        value = filters.get(field)
        if value:
            model_field = field if field != "team" else "department"
            query = query.filter(getattr(Asset, model_field) == value)

    rows = []
    for asset in query.limit(int(filters.get("limit", 1000))).all():
        rows.append(
            _sanitize_row(
                {
                    "asset_id": asset.id,
                    "asset_tag": asset.asset_tag,
                    "name": asset.name,
                    "asset_type": asset.asset_type,
                    "status": asset.status,
                    "department": asset.department,
                    "location": asset.location,
                    "assigned_to_email": asset.assigned_to.email if asset.assigned_to else None,
                    "created_at": asset.created_at.isoformat() if asset.created_at else None,
                }
            )
        )
    return rows


def query_deployments(actor, filters=None):
    """Query deployments rows (basic implementation)."""
    filters = filters or {}
    query = Deployment.query.order_by(Deployment.created_at.desc())

    for field in ("environment", "status", "service_name"):
        value = filters.get(field)
        if value:
            query = query.filter(getattr(Deployment, field) == value)

    rows = []
    for deployment in query.limit(int(filters.get("limit", 1000))).all():
        rows.append(
            _sanitize_row(
                {
                    "deployment_id": deployment.id,
                    "service_name": deployment.service_name,
                    "version": deployment.version,
                    "environment": deployment.environment,
                    "status": deployment.status,
                    "is_rollback": deployment.is_rollback,
                    "deployer_email": deployment.deployer.email if deployment.deployer else None,
                    "deployed_at": deployment.deployed_at.isoformat() if deployment.deployed_at else None,
                }
            )
        )
    return rows


def _project_rows(rows, fields):
    projected = []
    for row in rows:
        projected.append({field: row.get(field) for field in fields})
    return projected


def resolve_requested_fields(filters, template_allowed_fields):
    requested_fields = (filters or {}).get("fields")
    if requested_fields is None:
        return list(template_allowed_fields)

    if not isinstance(requested_fields, list) or any(not isinstance(field, str) for field in requested_fields):
        raise ValueError("filters_json.fields must be an array of field names")

    invalid = sorted(set(requested_fields) - set(template_allowed_fields))
    if invalid:
        raise ValueError(f"Requested fields are not allowed by template: {', '.join(invalid)}")

    return requested_fields


def apply_masking(rows, masking_rules):
    masking_rules = masking_rules or {}
    if not masking_rules:
        return rows

    masked_rows = []
    for row in rows:
        copy_row = dict(row)
        for field, rule in masking_rules.items():
            if field not in copy_row:
                continue
            value = copy_row.get(field)
            if value in (None, ""):
                continue
            if rule == "full":
                copy_row[field] = "***masked***"
            elif rule == "partial":
                value_str = str(value)
                copy_row[field] = ("*" * max(0, len(value_str) - 4)) + value_str[-4:]
        masked_rows.append(copy_row)
    return masked_rows


def generate_csv(rows, template_id):
    output = io.StringIO()
    if not rows:
        return "".encode("utf-8")

    fieldnames = list(rows[0].keys())
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()
    writer.writerows(rows)
    return output.getvalue().encode("utf-8")


def generate_markdown(rows, template_id):
    if not rows:
        return "".encode("utf-8")

    headers = list(rows[0].keys())
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join([str(row.get(h, "") or "") for h in headers]) + " |")
    return "\n".join(lines).encode("utf-8")


def generate_xlsx(rows, template_id):
    try:
        from openpyxl import Workbook
    except ImportError as exc:
        raise ValueError("XLSX generation requires openpyxl") from exc

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Export"

    if rows:
        headers = list(rows[0].keys())
        sheet.append(headers)
        for row in rows:
            sheet.append([row.get(header) for header in headers])

    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def generate_doc(template_id, variables=None, tables=None):
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("DOCX generation requires python-docx") from exc

    variables = variables or {}
    tables = tables or []

    doc = Document()
    doc.add_heading(variables.get("title") or f"ControlHub Report: {template_id}", level=1)

    subtitle = variables.get("subtitle")
    if subtitle:
        doc.add_paragraph(subtitle)

    for key, value in sorted(variables.items()):
        if key in {"title", "subtitle"}:
            continue
        doc.add_paragraph(f"{key}: {value}")

    for table_data in tables:
        table_title = table_data.get("title")
        if table_title:
            doc.add_heading(table_title, level=2)

        rows = table_data.get("rows") or []
        if not rows:
            doc.add_paragraph("No records.")
            continue

        headers = list(rows[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = str(header)

        for row in rows:
            cells = table.add_row().cells
            for idx, header in enumerate(headers):
                cells[idx].text = str(row.get(header, "") or "")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def sha256_hex(payload_bytes):
    return hashlib.sha256(payload_bytes).hexdigest()


def _local_artifact_path(s3_key):
    directory = os.environ.get("AGENT_ARTIFACTS_DIR", "/tmp/controlhub-agent-artifacts")
    os.makedirs(directory, mode=0o700, exist_ok=True)
    safe_key = s3_key.replace("/", "_")
    return os.path.join(directory, safe_key)


def store_artifact(file_bytes, artifact_id, filename, mime_type):
    """Store bytes and return bucket/key + expiry metadata."""
    backend = os.environ.get("AGENT_ARTIFACT_STORAGE", "local").lower()
    ttl_seconds = int(os.environ.get("AGENT_ARTIFACT_URL_EXPIRY_SECONDS", "300"))
    expires_at = datetime.utcnow().timestamp() + ttl_seconds
    expires_at_dt = datetime.utcfromtimestamp(expires_at)

    key = f"agent/{datetime.utcnow().strftime('%Y/%m/%d')}/{artifact_id}_{filename}"

    if backend == "s3":
        bucket = put_artifact_object(key, file_bytes, mime_type)
        return {
            "s3_bucket": bucket,
            "s3_key": key,
            "expires_at": expires_at_dt,
        }

    path = _local_artifact_path(key)
    with open(path, "wb") as handle:
        handle.write(file_bytes)
    os.chmod(path, 0o600)

    return {
        "s3_bucket": LOCAL_ARTIFACT_BUCKET,
        "s3_key": key,
        "expires_at": expires_at_dt,
    }


def read_artifact_bytes(s3_bucket, s3_key):
    if s3_bucket == LOCAL_ARTIFACT_BUCKET:
        path = _local_artifact_path(s3_key)
        with open(path, "rb") as handle:
            return handle.read()
    return read_artifact_object(s3_bucket, s3_key)


def generate_presigned_download(s3_bucket, s3_key, filename, expires_seconds):
    if s3_bucket == LOCAL_ARTIFACT_BUCKET:
        return None
    return generate_artifact_presigned_url(s3_bucket, s3_key, filename, expires_seconds)


def publish_to_drive(artifact_bytes, artifact, destination):
    folder_id = (destination.config or {}).get("folder_id")
    if not folder_id:
        raise ValueError("google_drive_folder destination requires folder_id")
    return google_publish_to_drive(
        file_bytes=artifact_bytes,
        filename=artifact.filename,
        mime_type=artifact.mime_type,
        folder_id=folder_id,
        subject_email=(
            os.environ.get("GOOGLE_IMPERSONATE_USER")
            or os.environ.get("GOOGLE_IMPERSONATED_USER")
        ),
    )


def publish_to_sheet(artifact_bytes, artifact, destination, mode="overwrite"):
    config = destination.config or {}
    spreadsheet_id = config.get("spreadsheet_id")
    sheet_name = config.get("sheet_name")
    a1_range = config.get("a1_range")

    if not spreadsheet_id or not a1_range:
        raise ValueError("google_sheet_range destination requires spreadsheet_id and a1_range")

    return google_publish_to_sheet(
        file_bytes=artifact_bytes,
        mime_type=artifact.mime_type,
        spreadsheet_id=spreadsheet_id,
        sheet_name=sheet_name,
        a1_range=a1_range,
        mode=mode,
        subject_email=(
            os.environ.get("GOOGLE_IMPERSONATE_USER")
            or os.environ.get("GOOGLE_IMPERSONATED_USER")
        ),
    )


def update_spreadsheet(sheet_id, cell_range, values):
    """Compatibility helper retained as an optional tool stub."""
    allowlist_raw = os.environ.get("AGENT_EXTERNAL_SHEET_ALLOWLIST", "")
    allowlist = {entry.strip() for entry in allowlist_raw.split(",") if entry.strip()}

    if sheet_id not in allowlist:
        raise ValueError("sheet_id is not allow-listed for external writes")

    row_count = len(values or [])
    col_count = len(values[0]) if row_count else 0
    return {
        "sheet_id": sheet_id,
        "range": cell_range,
        "updated_rows": row_count,
        "updated_columns": col_count,
        "status": "stubbed_success",
    }


def query_module_rows(actor, module_scope, filters):
    if module_scope == "people":
        return query_people(actor, filters)
    if module_scope == "assets":
        return query_assets(actor, filters)
    if module_scope == "deployments":
        return query_deployments(actor, filters)
    raise ValueError(f"Unsupported module_scope: {module_scope}")


def enforce_template_fields(rows, allowed_fields, selected_fields=None):
    if not allowed_fields:
        raise ValueError("Template contains no allowed fields")
    fields = selected_fields or list(allowed_fields)
    return _project_rows(rows, fields)


def to_json_serializable(data):
    return json.loads(json.dumps(data, default=str))
